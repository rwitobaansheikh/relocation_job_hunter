"""Render CVs and cover letters as Word .docx matching the CV360 ATS-friendly layout."""

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

# Calibri parses cleanly in Word and most ATS systems (matches CV360 exports).
_FONT = "Calibri"
_RIGHT_TAB = Inches(7.0)

_DEFAULT_SECTION_TITLES = {
    "summary": "PROFESSIONAL SUMMARY",
    "skills": "TECHNICAL SKILLS",
    "soft_skills": "SOFT SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "projects": "PROJECTS",
    "certifications": "CERTIFICATIONS & CONTINUOUS LEARNING",
    "education": "EDUCATION",
}


def role_abbrev(job_title: str) -> str:
    words = re.findall(r"[A-Za-z]+", job_title or "")
    if not words:
        return "Role"
    if len(words) == 1:
        return words[0][:6]
    return "".join(word[0].upper() for word in words[:4])


def document_filename(full_name: str, job_title: str, label: str) -> str:
    parts = [p for p in re.split(r"\s+", (full_name or "").strip()) if p]
    first = parts[0] if parts else "Candidate"
    last = parts[-1] if len(parts) > 1 else ""
    abbrev = role_abbrev(job_title)
    if last and last != first:
        base = f"{first}-{last}-{abbrev}-{label}"
    else:
        base = f"{first}-{abbrev}-{label}"
    cleaned = re.sub(r"[^\w\-.]", "", base.replace(" ", "-"))
    return f"{cleaned}.docx"


def cv_layout_prompt_reference() -> str:
    """Layout spec for the LLM — mirrors CV360-approved ATS structure."""
    return """
CV360 ATS LAYOUT (Word renderer reproduces this exactly):

HEADER (centred):
- Line 1: candidate full name (large, bold).
- Line 2: tailored professional tagline for the target role (bold, e.g. "Software Engineer and Machine Learning Specialist").
- Line 3: phone   |   email   |   linkedin   |   github — spaced pipes, no https:// prefixes.

SECTION ORDER (always use this order in the "sections" array):
1. summary — title "PROFESSIONAL SUMMARY"
2. skills — title "TECHNICAL SKILLS"
3. soft_skills — title "SOFT SKILLS" (role-tailored interpersonal strengths)
4. experience — title "PROFESSIONAL EXPERIENCE"
5. projects — title "PROJECTS"
6. certifications — title "CERTIFICATIONS & CONTINUOUS LEARNING" (omit section if none)
7. education — title "EDUCATION" (always last)

SECTION HEADERS: ALL CAPS, bold, no underline rule.

PROFESSIONAL SUMMARY: one paragraph (3-5 sentences), no bullets. Mention years of experience,
core domain, and 2-3 keywords from the job description naturally.

TECHNICAL SKILLS: 3-5 groups. Each group is one line:
  Category label (bold): comma-separated tools/technologies.
  Categories examples: "Machine Learning & Deep Learning", "MLOps, Cloud & DevOps", "Software & Data Engineering".
  Prioritise skills mentioned in the job description first.

SOFT SKILLS (critical — tailor to the target role): 4-6 groups, same one-line format as technical skills.
  Derive from the job description (e.g. leadership, communication, stakeholder management, adaptability).
  Ground each line in REAL evidence from the source CV — never invent experience.
  Examples:
  - Leadership & Team Development: led shift teams, mentored new starters, …
  - Communication: translated technical findings for non-technical stakeholders, …
  - Organised & Methodical: managed SLAs, triaged high-volume tickets, …

PROFESSIONAL EXPERIENCE (reverse chronological):
- Each entry: role and company combined as "Role — Company" (use em dash —).
- Date range and location on the same line, right-aligned: "Jul 2024 – Present | Bath, UK".
- 2-4 achievement bullets with metrics where possible. Store bullets WITHOUT a leading dash.

PROJECTS:
- Line 1: project name (bold).
- Line 2: tech stack | year (e.g. "PyTorch, Flask | 2026").
- 2-3 impact bullets.
- links: [{"label": "Source code", "url": "..."}] — only from ORIGINAL PROJECT LINKS.

CERTIFICATIONS: items with heading (cert/course name + provider), date, optional detail text.

EDUCATION (reverse chronological):
- Line 1: "Degree (grade if any), Institution" … date range on the right.
- Line 2: location (city, country).
- Optional "Relevant courses:" sentence listing modules aligned to the target role.

DATE FORMAT: Month abbreviations (Jan, Feb, Mar, Apr, May, Jun, Jul, Aug, Sep, Oct, Nov, Dec).
Use en dash – in displayed date ranges (Jul 2024 – Present).

Do NOT use tables, text boxes, columns, or graphics. Plain text only for ATS parsing.
"""


def _normalize_dates(text: str) -> str:
    return (text or "").replace("—", "–").replace("-", "–").replace("  ", " ")


def _display_url(url: str) -> str:
    label = (url or "").strip()
    for prefix in ("https://", "http://", "www."):
        if label.lower().startswith(prefix):
            label = label[len(prefix):]
    return label.rstrip("/")


def _set_run_font(run, size_pt: float = 10, bold: bool = False, italic: bool = False) -> None:
    run.font.name = _FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)


def _add_section_heading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title.upper())
    _set_run_font(run, size_pt=12, bold=True)


def _add_skill_line(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.left_indent = Inches(0.05)
    if label:
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, bold=True)
    if value:
        value_run = paragraph.add_run(value)
        _set_run_font(value_run)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    run = paragraph.add_run(text.strip())
    _set_run_font(run, size_pt=10)


def _add_experience_header(doc: Document, item: dict) -> None:
    """Role — Company on the left; date | location tab-aligned on the right."""
    role = (item.get("role") or "").strip()
    company = (item.get("company") or "").strip()
    heading = (item.get("heading") or "").strip()
    subheading = (item.get("subheading") or "").strip()

    if role and company:
        left = f"{role} — {company}"
    elif heading and subheading and "—" not in heading:
        left = f"{subheading} — {heading}"
    elif heading:
        left = heading
    else:
        left = role or company or subheading

    date = _normalize_dates(item.get("date") or "").strip()
    location = (item.get("location") or "").strip()
    right = f"{date} | {location}".strip(" |") if date or location else ""

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(2)
    left_run = paragraph.add_run(left)
    _set_run_font(left_run, bold=True)
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right)
        _set_run_font(right_run, size_pt=9)


def _add_education_header(doc: Document, item: dict) -> None:
    degree = (item.get("degree") or item.get("subheading") or "").strip()
    school = (item.get("school") or item.get("heading") or "").strip()
    if degree and school and degree not in school:
        left = f"{degree}, {school}"
    else:
        left = school or degree

    date = _normalize_dates(item.get("date") or "").strip()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(1)
    left_run = paragraph.add_run(left)
    _set_run_font(left_run, bold=True)
    if date:
        paragraph.add_run("\t")
        date_run = paragraph.add_run(date)
        _set_run_font(date_run, size_pt=9)

    location = (item.get("location") or "").strip()
    if location:
        loc_p = doc.add_paragraph(location)
        loc_p.paragraph_format.space_after = Pt(2)
        for run in loc_p.runs:
            _set_run_font(run, size_pt=9)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(colour)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _render_skill_groups(doc: Document, groups: list) -> None:
    for group in groups:
        if not isinstance(group, dict):
            continue
        label = (group.get("label") or "").strip()
        value = (group.get("value") or "").strip()
        if label or value:
            _add_skill_line(doc, label, value)


def _plain_paragraphs_from_text(doc: Document, text: str) -> None:
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text or "") if b.strip()]
    if not blocks and (text or "").strip():
        blocks = [line.strip() for line in text.splitlines() if line.strip()]
    for block in blocks:
        paragraph = doc.add_paragraph(block)
        for run in paragraph.runs:
            _set_run_font(run)


def render_resume_docx(data: dict, out_path: str) -> bool:
    """Render structured resume JSON to Word — CV360 ATS layout."""
    try:
        doc = Document()
        _configure_page(doc)

        name = (data.get("name") or "").strip()
        if name:
            name_p = doc.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_p.add_run(name)
            _set_run_font(name_run, size_pt=20, bold=True)

        tagline = (data.get("tagline") or "").strip()
        if tagline:
            tag_p = doc.add_paragraph()
            tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tag_run = tag_p.add_run(tagline)
            _set_run_font(tag_run, size_pt=11, bold=True)

        contact = data.get("contact") or {}
        contact_bits: list[str] = []
        for key in ("phone", "email", "linkedin", "github"):
            raw = (contact.get(key) or "").strip()
            if not raw:
                continue
            if key in ("linkedin", "github"):
                contact_bits.append(_display_url(raw))
            else:
                contact_bits.append(raw)
        if contact_bits:
            contact_p = doc.add_paragraph("   |   ".join(contact_bits))
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_p.paragraph_format.space_after = Pt(6)
            for run in contact_p.runs:
                _set_run_font(run, size_pt=9)

        for section in data.get("sections") or []:
            if not isinstance(section, dict):
                continue
            section_type = section.get("type") or ""
            title = (section.get("title") or "").strip() or _DEFAULT_SECTION_TITLES.get(
                section_type, section_type.upper()
            )
            _add_section_heading(doc, title)

            if section_type == "summary":
                text = (section.get("text") or "").strip()
                if text:
                    summary_p = doc.add_paragraph(text)
                    summary_p.paragraph_format.space_after = Pt(4)
                    for run in summary_p.runs:
                        _set_run_font(run)

            elif section_type in ("skills", "soft_skills"):
                _render_skill_groups(doc, section.get("groups") or [])

            elif section_type == "experience":
                for item in section.get("items") or []:
                    if not isinstance(item, dict):
                        continue
                    _add_experience_header(doc, item)
                    for bullet in item.get("bullets") or []:
                        if bullet:
                            _add_bullet(doc, str(bullet))

            elif section_type == "projects":
                for item in section.get("items") or []:
                    if not isinstance(item, dict):
                        continue
                    project_name = (item.get("heading") or "").strip()
                    tech = (item.get("tech") or "").strip()
                    date = _normalize_dates(item.get("date") or "").strip()

                    if project_name:
                        name_p = doc.add_paragraph()
                        name_p.paragraph_format.space_after = Pt(1)
                        name_run = name_p.add_run(project_name)
                        _set_run_font(name_run, bold=True)

                    if tech or date:
                        meta = " | ".join(x for x in (tech, date) if x)
                        meta_p = doc.add_paragraph(meta)
                        meta_p.paragraph_format.space_after = Pt(2)
                        for run in meta_p.runs:
                            _set_run_font(run, size_pt=9)

                    for bullet in item.get("bullets") or []:
                        if bullet:
                            _add_bullet(doc, str(bullet))

                    links = item.get("links") or []
                    if links:
                        links_p = doc.add_paragraph()
                        links_p.paragraph_format.space_after = Pt(6)
                        for idx, link in enumerate(links):
                            if not isinstance(link, dict):
                                continue
                            label = (link.get("label") or "Link").strip()
                            url = (link.get("url") or "").strip()
                            if idx > 0:
                                sep_run = links_p.add_run("   |   ")
                                _set_run_font(sep_run, size_pt=9)
                            _add_hyperlink(links_p, label, url)

            elif section_type == "certifications":
                for item in section.get("items") or []:
                    if not isinstance(item, dict):
                        continue
                    heading = (item.get("heading") or "").strip()
                    detail = (item.get("detail") or item.get("text") or "").strip()
                    if heading:
                        _add_bullet(doc, heading if not detail else f"{heading} — {detail}")
                    elif detail:
                        _add_bullet(doc, detail)

            elif section_type == "education":
                for item in section.get("items") or []:
                    if not isinstance(item, dict):
                        continue
                    _add_education_header(doc, item)
                    courses = (item.get("courses") or item.get("detail") or "").strip()
                    if courses:
                        if not courses.lower().startswith("relevant"):
                            courses = f"Relevant courses: {courses}"
                        courses_p = doc.add_paragraph(courses)
                        courses_p.paragraph_format.space_after = Pt(4)
                        for run in courses_p.runs:
                            _set_run_font(run, size_pt=9.5)

        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return True
    except Exception as exc:
        logger.error("Resume DOCX render failed: %s", exc)
        return False


def render_cover_letter_docx(name: str, contact: dict, body_text: str, out_path: str) -> bool:
    try:
        doc = Document()
        _configure_page(doc)

        if name:
            header = doc.add_paragraph()
            header_run = header.add_run(name)
            _set_run_font(header_run, size_pt=14, bold=True)

        contact_bits: list[str] = []
        for key in ("email", "phone", "linkedin", "location"):
            value = (contact.get(key) or "").strip()
            if value:
                contact_bits.append(value)
        if contact_bits:
            contact_p = doc.add_paragraph(" | ".join(contact_bits))
            for run in contact_p.runs:
                _set_run_font(run)

        cleaned = re.sub(
            r"^```[a-zA-Z]*\s*\n(.*?)\n?```$",
            r"\1",
            (body_text or "").strip(),
            flags=re.DOTALL,
        )
        _plain_paragraphs_from_text(doc, cleaned.strip())

        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return True
    except Exception as exc:
        logger.error("Cover letter DOCX render failed: %s", exc)
        return False
